import os
import re
import json
import logging
import platform
import shutil
import subprocess
import zipfile
import xml.etree.ElementTree as ET
import datetime
import tempfile
import threading
import time
from docxtpl import DocxTemplate
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    import mammoth
except ImportError:
    mammoth = None

logger = logging.getLogger(__name__)

_JINJA_TAG_NORMALIZE_PATTERNS = [
    (re.compile(r"\{\%\s*end\s+for\s*\%\}", re.IGNORECASE), "{% endfor %}"),
    (re.compile(r"\{\%\s*end\s+if\s*\%\}", re.IGNORECASE), "{% endif %}"),
    (re.compile(r"\{\%\s*end\s+block\s*\%\}", re.IGNORECASE), "{% endblock %}"),
]


def get_preview_unavailable_reason():
    if mammoth is None:
        return "No se pudo mostrar la vista previa. Instala mammoth con: pip install mammoth"
    return None


def _normalize_jinja_in_docx(template_path):
    """
    Crea una copia temporal del DOCX corrigiendo variantes comunes de tags Jinja
    que Word suele fragmentar o modificar (por ejemplo: {% end for %}).
    """
    try:
        temp_dir = tempfile.mkdtemp(prefix="docx_tpl_")
        normalized_path = os.path.join(temp_dir, os.path.basename(template_path))

        with zipfile.ZipFile(template_path, "r") as src_zip, zipfile.ZipFile(normalized_path, "w") as dst_zip:
            for item in src_zip.infolist():
                raw = src_zip.read(item.filename)
                if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                    xml_text = raw.decode("utf-8", errors="ignore")
                    # Normaliza espacios no separables que rompen el parser de jinja.
                    xml_text = xml_text.replace("\u00a0", " ")
                    # Quita marcas de ortografía/gramática que Word inserta en medio de tags.
                    xml_text = re.sub(r"<w:proofErr[^>]*/>", "", xml_text)
                    for pattern, replacement in _JINJA_TAG_NORMALIZE_PATTERNS:
                        xml_text = pattern.sub(replacement, xml_text)
                    raw = xml_text.encode("utf-8")

                dst_zip.writestr(item, raw)

        return normalized_path
    except Exception as exc:
        logger.warning(f"No se pudo normalizar plantilla DOCX ({template_path}): {exc}")
        return template_path

# Configuración y estilos para las tablas extraidas
TABLE_STYLES_CONFIG = {
    "font_size": 9,
    "font_name": "Arial",
    "header_bg_color": "D9EAF7",
    "header_font_color": "000000",
    "header_font_bold": True,
    "border_color": "000000",
    "border_size": 4
}


def _set_cell_shading(cell, fill_color_hex):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(fill_color_hex or "D9EAF7"))
    tc_pr.append(shd)

def extract_variables_from_template(file_path):
    """
    Lee un documento Word (docx) y extrae todas las variables definidas como {{ variable }}
    Usando el motor de docxtpl.
    """
    pattern = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_\.-]*)\s*\}\}")

    internal_vars = {"tabla_items", "tabla_columnas", "tabla_filas", "celda"}

    def _normalize(found):
        normalized = []
        for item in found:
            name = str(item or "").strip().lower()
            if not name:
                continue
            if name in internal_vars:
                continue
            if name.startswith("item.") or name.startswith("col.") or name.startswith("fila."):
                continue
            if name not in normalized:
                normalized.append(name)
        return normalized

    def _extract_with_docxtpl(path):
        doc = DocxTemplate(path)
        variables = doc.get_undeclared_template_variables() or set()
        return _normalize(list(variables))

    def _extract_with_python_docx(path):
        found = set()
        doc = Document(path)

        for p in doc.paragraphs:
            for match in pattern.findall(p.text or ""):
                found.add(match)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for match in pattern.findall(cell.text or ""):
                        found.add(match)

        for section in doc.sections:
            for p in section.header.paragraphs:
                for match in pattern.findall(p.text or ""):
                    found.add(match)
            for p in section.footer.paragraphs:
                for match in pattern.findall(p.text or ""):
                    found.add(match)

        return _normalize(list(found))

    def _extract_with_xml(path):
        found = set()
        with zipfile.ZipFile(path, "r") as zip_docx:
            targets = [
                name for name in zip_docx.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            ]
            for name in targets:
                try:
                    xml = zip_docx.read(name).decode("utf-8", errors="ignore")
                except Exception:
                    continue
                try:
                    root = ET.fromstring(xml)
                    text_nodes = [node.text or "" for node in root.iter() if node.tag.endswith("}t")]
                    xml_text = "\n".join(text_nodes)
                except Exception:
                    xml_text = re.sub(r"<[^>]+>", "", xml)

                for match in pattern.findall(xml_text):
                    found.add(match)
        return _normalize(list(found))

    try:
        variables = _extract_with_docxtpl(file_path)
        if variables:
            return variables
    except Exception as e:
        logger.info(f"Extractor docxtpl no pudo parsear {file_path}; usando fallback: {e}")

    try:
        variables = _extract_with_python_docx(file_path)
        if variables:
            return variables
    except Exception as e:
        logger.warning(f"Extractor python-docx falló para {file_path}: {e}")

    try:
        return _extract_with_xml(file_path)
    except Exception as e:
        logger.error(f"Error extrayendo variables de {file_path}: {e}")
        return []

def configure_table_data(extraction_data):
    if not extraction_data:
        return []

    if not isinstance(extraction_data, list):
        return []

    normalized_rows = []
    for raw_row in extraction_data:
        if not isinstance(raw_row, dict):
            continue

        normalized_row = {}
        for key, value in raw_row.items():
            normalized_key = str(key).strip()
            if not normalized_key:
                continue

            if value is None:
                normalized_row[normalized_key] = ""
            elif isinstance(value, bool):
                normalized_row[normalized_key] = "Si" if value else "No"
            elif isinstance(value, float):
                normalized_row[normalized_key] = f"{value:.2f}".rstrip("0").rstrip(".")
            else:
                text_value = str(value).strip()
                if text_value.lower() in {"none", "null", "nan", "undefined"}:
                    text_value = ""
                normalized_row[normalized_key] = text_value

        if normalized_row:
            normalized_rows.append(normalized_row)

    return normalized_rows

def build_dynamic_table_context(table_rows, table_columns):
    if not table_rows or not table_columns:
        return [], []

    columnas = []
    for col in table_columns:
        col_id = str(col.get("id", "")).strip()
        if not col_id:
            continue
        columnas.append({
            "id": col_id,
            "label": str(col.get("label", col_id)).strip() or col_id,
        })

    if not columnas:
        return [], []

    filas = []
    for row in table_rows:
        celdas = []
        for col in columnas:
            raw_value = row.get(col["id"], "-")
            if raw_value is None:
                celdas.append("-")
            elif isinstance(raw_value, bool):
                celdas.append("Si" if raw_value else "No")
            else:
                text = str(raw_value).strip()
                if text.lower() in {"none", "null", "nan", "undefined"}:
                    text = ""
                celdas.append(text if text else "-")
        filas.append({"celdas": celdas})

    return columnas, filas


def _set_cell_text(cell, value, *, bold=False, align=WD_PARAGRAPH_ALIGNMENT.LEFT, font_size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(value if value is not None else ""))
    run.bold = bool(bold)
    run.font.size = Pt(int(font_size))


def _build_area_summary_table_subdoc(subdoc, filas, context_data):
    table = subdoc.add_table(rows=1, cols=2)
    for style_name in ("Table Grid", "Tabla con cuadricula", "Tabla con cuadrícula"):
        try:
            table.style = style_name
            break
        except Exception:
            continue

    table.autofit = False

    titulo = str((context_data or {}).get("titulo_acta") or "").strip()
    if not titulo:
        area_nombre = str((context_data or {}).get("nombre_area") or "").strip()
        numero_acta = str((context_data or {}).get("numero_acta") or "").strip()
        if area_nombre and numero_acta:
            titulo = f"{area_nombre} ACTA No. {numero_acta}"
        else:
            titulo = "ACTA"

    header_cells = table.rows[0].cells
    merged_header = header_cells[0].merge(header_cells[1])
    _set_cell_text(
        merged_header,
        titulo,
        bold=True,
        align=WD_PARAGRAPH_ALIGNMENT.CENTER,
        font_size=11,
    )

    for fila in filas:
        row_cells = table.add_row().cells
        celdas = fila.get("celdas", [])
        descripcion = str(celdas[0] if len(celdas) > 0 else "-")
        cantidad = str(celdas[1] if len(celdas) > 1 else "-")

        is_total = descripcion.strip().upper() == "TOTAL DE BIENES"
        align_desc = WD_PARAGRAPH_ALIGNMENT.CENTER if is_total else WD_PARAGRAPH_ALIGNMENT.LEFT
        align_cant = WD_PARAGRAPH_ALIGNMENT.CENTER

        _set_cell_text(row_cells[0], descripcion, bold=is_total, align=align_desc, font_size=9)
        _set_cell_text(row_cells[1], cantidad, bold=is_total, align=align_cant, font_size=9)

    desc_values = []
    for fila in filas:
        celdas = fila.get("celdas", [])
        if not celdas:
            continue
        descripcion = str(celdas[0] or "").strip()
        if descripcion.upper() == "TOTAL DE BIENES":
            continue
        desc_values.append(descripcion)

    max_desc_len = max((len(v) for v in desc_values), default=0)
    use_compact_size = max_desc_len > 0 and max_desc_len < 15

    desc_width = 5.9
    qty_width = 1.1
    if use_compact_size:
        desc_width = round(desc_width * 0.64, 2)
        qty_width = round(qty_width * 0.64, 2)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER if use_compact_size else WD_TABLE_ALIGNMENT.LEFT

    for row in table.rows:
        row.cells[0].width = Inches(desc_width)
        row.cells[1].width = Inches(qty_width)

    return subdoc


def _compute_general_table_widths(columnas, total_width_in=7.0):
    cols = list(columnas or [])
    n = len(cols)
    if n <= 0:
        return []

    base = total_width_in / n
    widths = [base for _ in range(n)]

    desc_idx = -1
    estado_idx = -1
    for idx, col in enumerate(cols):
        col_id = str(col.get("id") or "").strip().lower()
        col_label = str(col.get("label") or "").strip().lower()
        if desc_idx < 0 and (col_id == "descripcion" or "descrip" in col_label):
            desc_idx = idx
        if estado_idx < 0 and (col_id == "estado" or "estado" in col_label):
            estado_idx = idx

    if estado_idx >= 0 and desc_idx >= 0 and estado_idx != desc_idx:
        old_estado = widths[estado_idx]
        new_estado = max(0.9, round(old_estado - 0.35, 2))
        delta_estado = old_estado - new_estado
        widths[estado_idx] = new_estado
        widths[desc_idx] = round(widths[desc_idx] + delta_estado, 2)

    diff = round(total_width_in - sum(widths), 4)
    if abs(diff) > 0 and n > 0:
        target_idx = desc_idx if desc_idx >= 0 else 0
        widths[target_idx] = round(widths[target_idx] + diff, 2)

    return widths


def build_dynamic_table_subdoc(doc_tpl, table_rows, table_columns, context_data=None):
    if not table_rows or not table_columns:
        return None

    columnas, filas = build_dynamic_table_context(table_rows, table_columns)
    if not columnas:
        return None

    try:
        subdoc = doc_tpl.new_subdoc()
    except Exception as exc:
        if "docxcompose" in str(exc).lower():
            logger.warning("docxcompose no instalado. tabla_dinamica se omitira temporalmente.")
            return None
        raise
    normalized_ids = [str(col.get("id", "")).strip().lower() for col in columnas]
    if len(columnas) == 2 and normalized_ids == ["descripcion", "cantidad"]:
        return _build_area_summary_table_subdoc(subdoc, filas, context_data)

    table = subdoc.add_table(rows=1, cols=len(columnas))
    for style_name in ("Table Grid", "Tabla con cuadricula", "Tabla con cuadrícula"):
        try:
            table.style = style_name
            break
        except Exception:
            continue
    table.autofit = False

    column_widths = _compute_general_table_widths(columnas, total_width_in=7.0)

    # Encabezados
    header_cells = table.rows[0].cells
    for idx, col in enumerate(columnas):
        _set_cell_shading(header_cells[idx], TABLE_STYLES_CONFIG.get("header_bg_color", "D9EAF7"))
        run = header_cells[idx].paragraphs[0].add_run(str(col.get("label", "")))
        run.bold = bool(TABLE_STYLES_CONFIG.get("header_font_bold", True))
        run.font.size = Pt(int(TABLE_STYLES_CONFIG.get("font_size", 10)))

    # Filas de datos
    for fila in filas:
        row_cells = table.add_row().cells
        celdas = fila.get("celdas", [])
        for idx in range(len(columnas)):
            val = celdas[idx] if idx < len(celdas) else "-"
            run = row_cells[idx].paragraphs[0].add_run(str(val if val is not None else "-"))
            run.font.size = Pt(int(TABLE_STYLES_CONFIG.get("font_size", 9)))

    for row in table.rows:
        for idx, w in enumerate(column_widths):
            row.cells[idx].width = Inches(w)

    return subdoc

def generate_acta(
    template_path,
    context_data,
    table_data=None,
    table_columns=None,
    output_dir=None,
    doc_name="acta",
    use_date_subfolder=True,
    include_time_suffix=True,
):
    """
    Genera únicamente el archivo DOCX final.
    """
    normalized_template = _normalize_jinja_in_docx(template_path)
    doc = DocxTemplate(normalized_template)
    render_context = dict(context_data or {})
    
    if table_data:
        render_context['tabla_items'] = configure_table_data(table_data)

    if table_data and table_columns:
        tabla_columnas, tabla_filas = build_dynamic_table_context(table_data, table_columns)
        render_context['tabla_columnas'] = tabla_columnas
        render_context['tabla_filas'] = tabla_filas
        render_context['tabla_dinamica'] = build_dynamic_table_subdoc(
            doc,
            table_data,
            table_columns,
            context_data=render_context,
        )
        
    try:
        doc.render(render_context)
    except Exception as e:
        msg = str(e)
        if "unknown tag 'endfor'" in msg.lower() or "unexpected 'endfor'" in msg.lower():
            raise ValueError(
                "Sintaxis de tabla dinámica inválida en Word. "
                "Recomendado: use una sola variable {{ tabla_dinamica }} en la celda donde debe ir la tabla. "
            )
        raise
    
    if not output_dir:
        home = os.path.expanduser("~")
        output_dir = os.path.join(home, "Downloads", "inventario", doc_name.lower())
        
    fecha_str = datetime.datetime.now().strftime("%d-%m-%Y")
    final_output_dir = os.path.join(output_dir, fecha_str) if use_date_subfolder else output_dir
    
    if not os.path.exists(final_output_dir):
        os.makedirs(final_output_dir, exist_ok=True)
    
    if include_time_suffix:
        time_str = datetime.datetime.now().strftime("%H%M%S")
        docx_filename = f"{doc_name}_{time_str}.docx"
    else:
        docx_filename = f"{doc_name}.docx"
    
    docx_path = os.path.join(final_output_dir, docx_filename)
    
    try:
        doc.save(docx_path)
    except Exception as e:
        logger.error(f"Error guardando DOCX en {docx_path}: {e}")
        return None
        
    return docx_path


def render_docx_preview_html(docx_path):
    """
    Convierte DOCX a HTML simple para vista previa.
    """
    if not docx_path or not os.path.exists(docx_path) or mammoth is None:
        return None

    try:
        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value or ""
        if not html.strip():
            return None

        return (
            "<html><head><meta charset='utf-8'></head>"
            "<body style='font-family: Arial, sans-serif; margin: 20px;'>"
            f"{html}"
            "</body></html>"
        )
    except Exception as exc:
        logger.warning(f"No se pudo generar preview HTML for {docx_path}: {exc}")
        return None
